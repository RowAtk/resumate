from resumate import nlp
from resumate.iengines.utils import *

# import and register all IEngines 
from resumate.iengines.education.ie_eduction import ieducation
from resumate.iengines.skill.ie_skill import ieskills
from resumate.iengines.project.ie_project import ieproject

# Doc Generator
from docx import Document


"Takes Knowledge from the Various IE and outputs them to a doc file"

def createDoc(filename, user):
    document = Document()
    nl = "\n" #for new line

    # if not user.exists():
    #     user.fullname = "Marke Clarke"
    #     user.address = "UWI Mona"
    #     user.email = "mclarke@gmail.com"
    #     user.telnum = "876-995-9656"

    document.add_heading(f'Resume for {user.fullname}', 0)
    # CONTACT INFO
    document.add_heading('Contact Info', 1)
    document.add_paragraph(f'Address: {user.address}', style = "List Bullet")
    document.add_paragraph(f'Email: {user.email}', style = "List Bullet")
    document.add_paragraph(f'Tel #: {user.telnum}', style = "List Bullet")

    # EDUCATION
    document.add_heading('Education', 1)

    degrees = ieducation.iobjects

    for d in degrees:
        
        # p = document.add_paragraph(str(d))
        try:
            document.add_paragraph(f'{d.properties["title"]}{nl}{d.properties["date"].strftime("%Y")}{nl}{d.properties["source"]}')
        except:
            pass
        print(str(d))

    # SKILLS
    document.add_heading('Skills', 1)

    skills = ieskills.iobjects[0]
    print(skills)
    for key, vals in skills.properties.items():
        if vals:
            p = document.add_paragraph('', style="List Bullet")
            p.add_run(f'{key.title()}: ').bold = True
            p.add_run(", ".join(set(vals))).italic = True
    

    # PROJECTS
    document.add_heading('Projects', 1)

    projects = ieproject.iobjects
    document.add_heading('Work in Progress', 2)
    for proj in projects:
        document.add_paragraph(str(proj))

    print(projects)
#     for proj in projects:
#         document.add_heading(f'{proj.properties["proj title"]}', 2)
#         document.add_paragraph(f'Date: {proj.properties["date proj"].strftime("%Y")}')
#         projhelp(document, proj.properties["role"], "Role")
#         projhelp(document, proj.properties["description"], "Description")
#         projhelp(document, proj.properties["achieve"], "Achievements")

    document.save(filename)

# def projhelp(document, prop, title):
#     document.add_heading(f'{title}:', 3)
#     for val in prop:
#         if val:
#             p = document.add_paragraph('', style="List Bullet")
#             p.add_run(f'{val}')
